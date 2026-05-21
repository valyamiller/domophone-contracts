# pyright: reportMissingImports=false
# ruff: noqa

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def generate_contract(client, amount=480):
    """Генерация договора из шаблона Word"""
    
    template_path = os.path.join(os.path.dirname(__file__), 'contract_template.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Файл шаблона не найден: {template_path}")
    
    doc = Document(template_path)
    
    current_date = datetime.now().strftime('%d.%m.%Y')
    start_date = client.contract_start.strftime('%d.%m.%Y')
    end_date = client.contract_end.strftime('%d.%m.%Y')
    
    # Полный адрес
    full_address = f"{client.microdistrict}-{client.house}-{client.apartment}"
    
    # Формируем пункт о монтаже только если трубка установлена
    installation_clause = ""
    if client.tube_installed:
        installation_clause = "3.2. Плата за монтаж абонентского устройства (АУ) составляет с одной квартиры, 3000 (три тысячи) рублей 00 коп."
    
    # Все замены
    replacements = {
        '{microdistrict}': client.microdistrict,
        '{house}': client.house,
        '{apartment}': client.apartment,
        '{full_name}': client.full_name,
        '{phone}': client.phone,
        '{current_date}': current_date,
        '{start_date}': start_date,
        '{end_date}': end_date,
        '{amount}': str(amount),
        '{amount_text}': 'четыреста восемьдесят',
        '{personal_account}': client.personal_account,
        '{contract_number}': str(client.contract_number),
        '{installation_clause}': installation_clause,
        '{full_address}': full_address,
    }
    
    # Создаём или получаем единый стиль для всего документа
    try:
        # Пробуем получить существующий стиль
        normal_style = doc.styles['Normal']
    except KeyError:
        # Создаём новый стиль
        normal_style = doc.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    
    # Настраиваем шрифт для стиля
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(14)
    
    # Заменяем текст во всех параграфах и применяем стиль
    for paragraph in doc.paragraphs:
        # Заменяем текст
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
        
        # Применяем единый стиль ко всем параграфам
        paragraph.style = normal_style
        
        # Для каждой части текста принудительно устанавливаем шрифт
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Заменяем текст
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
                    
                    # Применяем стиль
                    paragraph.style = normal_style
                    
                    # Устанавливаем шрифт
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
    
    # Если трубка НЕ установлена - перенумеровываем пункты
    if not client.tube_installed:
        # Словарь для замены номеров пунктов
        renumber_map = [
            ('3.3.', '3.2.'),
            ('3.4.', '3.3.'),
            ('3.5.', '3.4.'),
        ]
        
        for paragraph in doc.paragraphs:
            for old, new in renumber_map:
                if old in paragraph.text:
                    paragraph.text = paragraph.text.replace(old, new)
            
            # Снова применяем стиль после замены
            paragraph.style = normal_style
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
    
    # Сохраняем договор
    contracts_dir = os.path.join(os.path.dirname(__file__), 'contracts')
    if not os.path.exists(contracts_dir):
        os.makedirs(contracts_dir)
    
    filename = f"Договор_{client.contract_number}_{client.personal_account}.docx"
    filepath = os.path.join(contracts_dir, filename)
    
    doc.save(filepath)
    
    return filepath, filename